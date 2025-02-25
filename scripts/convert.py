#!/usr/bin/env python3
import os
import re
import sys
import argparse
import logging
import yaml
import fnmatch
import shutil
import zipfile
import docx2pdf  # type: ignore
import docx  # type: ignore
import xml.etree.ElementTree as ElTree
from typing import List, Dict, Tuple, Any
from operator import itemgetter
from itertools import groupby


class ConvertVars:
    BASE_PATH = os.path.split(os.path.dirname(os.path.realpath(__file__)))[0]
    FILETYPE_CHOICES: List[str] = ["all", "docx", "pdf", "idml"]
    LANGUAGE_CHOICES: List[str] = ["template", "all", "en", "es", "fr", "pt-br"]
    DEFAULT_TEMPLATE_FILENAME: str = os.sep.join(
        ["resources", "templates", "owasp_cornucopia_edition_lang_ver_template"]
    )
    DEFAULT_OUTPUT_FILENAME: str = os.sep.join(["output", "owasp_cornucopia_edition_component_lang_ver"])
    args: argparse.Namespace
    can_convert_to_pdf: bool = False
    making_template: bool = False


def main() -> None:
    convert_vars.args = parse_arguments(sys.argv[1:])
    set_logging()
    logging.debug(" --- args = " + str(convert_vars.args))

    set_can_convert_to_pdf()
    if (
        convert_vars.args.outputfiletype == "pdf"
        and not convert_vars.can_convert_to_pdf
        and not convert_vars.args.debug
    ):
        logging.error(
            "Cannot convert to pdf on this system. "
            "Pdf conversion is available on Windows and Mac, if MS Word is installed"
        )
        return

    set_making_template()

    # Create output files
    for file_type in get_valid_file_types():
        for language in get_valid_language_choices():
            convert_type_language(file_type, language)


def get_valid_file_types() -> List[str]:
    if not convert_vars.args.outputfiletype:
        file_type = os.path.splitext(os.path.basename(convert_vars.args.outputfile))[1].strip(".")
        if file_type in ("", None):
            file_type = "docx"
        return [file_type]
    if convert_vars.args.outputfiletype.lower() == "pdf":
        if convert_vars.can_convert_to_pdf:
            return ["pdf"]
        else:
            logging.error("PDF output selected but currently unable to output PDF on this OS.")
            return []
    if convert_vars.args.outputfiletype.lower() == "all":
        file_types = []
        for file_type in convert_vars.FILETYPE_CHOICES:
            if file_type != "all" and (file_type != "pdf" or convert_vars.can_convert_to_pdf):
                file_types.append(file_type)
        return file_types
    if convert_vars.args.outputfiletype.lower() in convert_vars.FILETYPE_CHOICES:
        return [convert_vars.args.outputfiletype.lower()]
    return []


def get_valid_language_choices() -> List[str]:
    languages = []
    if convert_vars.args.language.lower() == "all":
        for language in convert_vars.LANGUAGE_CHOICES:
            if language not in ("all", "template"):
                languages.append(language)
    elif convert_vars.args.language == "":
        languages.append("en")
    else:
        languages.append(convert_vars.args.language)
    return languages


def set_logging() -> None:
    logging.basicConfig(
        format="%(asctime)s %(filename)s | %(levelname)s | %(funcName)s | %(message)s",
    )
    if convert_vars.args.debug:
        logging.getLogger().setLevel(logging.DEBUG)
    else:
        logging.getLogger().setLevel(logging.INFO)


def convert_type_language(file_type: str, language: str = "en") -> None:
    # Get the list of available translation files
    yaml_files = get_files_from_of_type(os.sep.join([convert_vars.BASE_PATH, "source"]), "yaml")
    if not yaml_files:
        return

    # Get the language data from the correct language file (checks vars.args.language to select the correct file)
    language_data: Dict[str, Dict[str, str]] = get_replacement_data(yaml_files, "translation", language)

    # Get the dict of replacement data
    language_dict: Dict[str, str] = get_replacement_dict(language_data, False)

    # Get meta data from language data
    meta: Dict[str, str] = get_meta_data(language_data)

    mapping_dict: Dict[str, str] = get_mapping_dict(yaml_files)

    if convert_vars.making_template:
        language_dict = remove_short_keys(language_dict)

    template_doc: str = get_template_doc(file_type)

    if not language_data or not mapping_dict or not meta or not template_doc:
        return

    # Name output file with correct edition, component, language & version
    output_file: str = rename_output_file(file_type, meta)
    ensure_folder_exists(os.path.dirname(output_file))

    # Work with docx file (and maybe convert to pdf afterwards)
    if file_type in ("docx", "pdf"):
        # Get the input (template) document
        doc: docx.Document = get_docx_document(template_doc)
        if convert_vars.making_template:
            doc = replace_docx_inline_text(doc, language_dict)
            doc = replace_docx_inline_text(doc, mapping_dict)
        else:
            language_dict.update(mapping_dict)
            doc = replace_docx_inline_text(doc, language_dict)

        if file_type == "docx":
            doc.save(output_file)
        else:
            # If file type is pdf, then save a temp docx file, convert the docx to pdf
            temp_docx_file = os.sep.join([convert_vars.BASE_PATH, "output", "temp.docx"])
            save_docx_file(doc, temp_docx_file)
            output_file = convert_docx_to_pdf(temp_docx_file, output_file)

    elif file_type == "idml":
        language_dict.update(mapping_dict)
        save_idml_file(template_doc, language_dict, output_file)

    logging.info("New file saved: " + str(output_file))


def save_idml_file(template_doc: str, language_dict: Dict[str, str], output_file: str) -> None:
    # Get the output path and temp output path to put the temp xml files
    output_path = convert_vars.BASE_PATH + os.sep + "output"
    temp_output_path = output_path + os.sep + "temp"
    # Ensure the output folder and temp output folder exist
    ensure_folder_exists(temp_output_path)
    logging.debug(" --- temp_folder for extraction of xml files = " + str(temp_output_path))

    # Unzip source xml files and place in temp output folder
    with zipfile.ZipFile(template_doc) as idml_archive:
        idml_archive.extractall(temp_output_path)
        logging.debug(" --- namelist of first few files in archive = " + str(idml_archive.namelist()[:5]))

    xml_files = get_files_from_of_type(temp_output_path, "xml")
    # Only Stories files have content to update
    for file in fnmatch.filter(xml_files, "*Stories*Story*"):
        if os.path.getsize(file) == 0:
            continue
        replace_text_in_xml_file(file, language_dict)

    # Zip the files as an idml file in output folder
    logging.debug(" --- finished replacing text in xml files. Now zipping into idml file")
    zip_dir(temp_output_path, output_file)

    # If not debugging, delete temp folder and files
    if not convert_vars.args.debug and os.path.exists(temp_output_path):
        shutil.rmtree(temp_output_path, ignore_errors=True)


def save_docx_file(doc: docx.Document, output_file: str) -> None:
    ensure_folder_exists(os.path.dirname(output_file))
    doc.save(output_file)


def convert_docx_to_pdf(docx_filename: str, output_pdf_filename: str) -> str:
    logging.debug(f" --- docx_file = {docx_filename}\n--- starting pdf conversion now.")
    fail: bool = False
    msg: str = ""
    if convert_vars.can_convert_to_pdf:
        # Use docx2pdf for windows and mac with MS Word installed
        try:
            docx2pdf.convert(docx_filename, output_pdf_filename)
        except Exception as e:
            msg = f"\nConvert error: {e}"
            fail = True
    else:
        fail = True
    if fail:
        msg = (
            "Error. A temporary docx file was created in the output folder but cannot be converted "
            f"to pdf (yet) on operating system: {sys.platform}\n"
            "This does work on Windows and Mac with MS Word installed."
        ) + msg
        logging.warning(msg)
        return docx_filename

    # If not debugging then delete the temp file
    if not convert_vars.args.debug:
        os.remove(docx_filename)
    return output_pdf_filename


def get_mapping_dict(yaml_files: List[str]) -> Dict[str, str]:
    mapping_data: Dict[str, Dict[str, str]] = get_replacement_data(yaml_files, "mappings")
    if not mapping_data:
        return {}
    return get_replacement_dict(mapping_data, True)


def set_can_convert_to_pdf() -> bool:
    operating_system: str = sys.platform.lower()
    can_convert = operating_system.find("win") != -1 or operating_system.find("darwin") != -1
    convert_vars.can_convert_to_pdf = can_convert
    logging.debug(f" --- operating system = {operating_system}, can_convert_to_pdf = {convert_vars.can_convert_to_pdf}")
    return can_convert


def sort_keys_longest_to_shortest(replacement_dict: Dict[str, str]) -> List[Tuple[str, str]]:
    new_list = list((k, v) for k, v in replacement_dict.items())
    return sorted(new_list, key=lambda s: len(s[0]), reverse=True)


def replace_text_in_xml_file(filename: str, replacement_dict: Dict[str, str]) -> None:
    if convert_vars.making_template:
        replacement_values = sort_keys_longest_to_shortest(replacement_dict)
    else:
        replacement_values = list(replacement_dict.items())

    try:
        tree = ElTree.parse(filename)
    except ElTree.ParseError as e:
        logging.error(f" --- parsing xml file: {filename}. error = {e}")
        return

    all_content_elements = tree.findall(".//Content")

    found_element = False
    for el in [el for el in all_content_elements]:
        if el.text == "" or el.text is None:
            continue
        el_text = get_replacement_value_from_dict(el.text, replacement_values)
        if el_text:
            el.text = el_text
            found_element = True

    if found_element:
        with open(filename, "bw") as f:
            f.write(ElTree.tostring(tree.getroot(), encoding="utf-8"))


def get_replacement_value_from_dict(el_text: str, replacement_values: List[Tuple[str, str]]) -> str:
    for (k, v) in replacement_values:
        k2: str = k.replace("'", "’").strip()
        v2: str = v.strip()
        if el_text == k:
            return v
        elif el_text.strip() == k2:
            return v2
        elif el_text.lower() == k.lower():
            return v
        elif el_text.strip().lower() == k2.lower():
            return v2
        elif convert_vars.making_template:
            reg_str = "^(OWASP SCP|OWASP ASVS|OWASP AppSensor|CAPEC|SAFECODE)\u2028" + k.replace(" ", "").strip() + "$"
            value_name = v[9:-1].replace("_", " ").lower().strip()
            new_text_test = (
                el_text[: len(value_name)] + "\u2028" + el_text[len(value_name) + 1 :].replace(" ", "").strip()
            )
            if re.match(reg_str, new_text_test) and el_text.lower().startswith(value_name):
                return el_text[: len(value_name)] + "\u2028" + v
        else:
            el_new = get_replacement_mapping_value(k, v, el_text)
            if el_new:
                return el_new
    return ""


def get_replacement_mapping_value(k: str, v: str, el_text: str) -> str:
    reg_str: str = "^(OWASP SCP|OWASP ASVS|OWASP AppSensor|CAPEC|SAFECODE)\u2028" + k.replace("$", "\\$").strip() + "$"
    if re.match(reg_str, el_text.strip()):
        pretext = el_text[: el_text.find("\u2028")]
        v_new = v
        if len(v) + len(pretext) > 60:
            v_new = v.replace(", ", ",")
        if len(v_new) >= 34:
            v_split = v_new.find(",", 25 - len(pretext)) + 1
            el_new = pretext + "    " + v_new[:v_split] + "\u2028" + v_new[v_split:].strip()
            return el_new
        else:
            return el_text.replace(k, v)
    return ""


def remove_short_keys(replacement_dict: Dict[str, str], min_length: int = 8) -> Dict[str, str]:
    data2: Dict[str, str] = {}
    for key, value in replacement_dict.items():
        if len(key) >= min_length:
            data2[key] = value
    logging.debug(
        " --- Making template. Removed card_numbers. len replacement_dict = "
        f"{len(replacement_dict)}, len data2 = {len(data2)}"
    )
    return data2


def get_template_doc(file_type: str) -> str:
    template_doc: str
    args_input_file: str = convert_vars.args.inputfile
    source_file_ext = file_type.replace("pdf", "docx")  # Pdf output uses docx source file
    if args_input_file:
        # Input file was specified
        if os.path.isabs(args_input_file):
            template_doc = args_input_file
        elif os.path.isfile(convert_vars.BASE_PATH + os.sep + args_input_file):
            template_doc = os.path.normpath(convert_vars.BASE_PATH + os.sep + args_input_file)
        elif os.path.isfile(convert_vars.BASE_PATH + os.sep + args_input_file.replace(".." + os.sep, "")):
            template_doc = os.path.normpath(
                convert_vars.BASE_PATH + os.sep + args_input_file.replace(".." + os.sep, "")
            )
        elif args_input_file.find("..") == -1 and os.path.isfile(
            convert_vars.BASE_PATH + os.sep + ".." + os.sep + args_input_file
        ):
            template_doc = os.path.normpath(convert_vars.BASE_PATH + os.sep + ".." + os.sep + args_input_file)
        elif os.path.isfile(convert_vars.BASE_PATH + os.sep + args_input_file.replace("scripts" + os.sep, "")):
            template_doc = os.path.normpath(
                convert_vars.BASE_PATH + os.sep + args_input_file.replace("scripts" + os.sep, "")
            )
        else:
            template_doc = args_input_file
            logging.debug(f" --- Template_doc NOT found. Input File = {args_input_file}")
    else:
        # No input file specified - using defaults
        if convert_vars.making_template:
            template_doc = os.sep.join(
                [convert_vars.BASE_PATH, "resources", "originals", "owasp_cornucopia_en." + source_file_ext]
            )
        else:
            template_doc = os.path.normpath(
                convert_vars.BASE_PATH + os.sep + convert_vars.DEFAULT_TEMPLATE_FILENAME + "." + source_file_ext
            )

    template_doc = template_doc.replace("\\ ", " ")
    if os.path.isfile(template_doc):
        template_doc = check_fix_file_extension(template_doc, source_file_ext)
        logging.debug(f" --- Returning template_doc = {template_doc}")
    else:
        logging.error(f"Source file not found: {args_input_file}. Please ensure file exists and try again.")
        template_doc = ""
    return template_doc


def rename_output_file(file_type: str, meta: Dict[str, str]) -> str:
    """Rename output file replacing place-holders from meta dict (edition, component, language, version)."""
    args_output_file: str = convert_vars.args.outputfile
    logging.debug(f" --- args_output_file = {args_output_file}")
    if args_output_file:
        # Output file is specified as an argument
        if os.path.isabs(args_output_file):
            output_filename = args_output_file
        else:
            output_filename = os.path.normpath(convert_vars.BASE_PATH + os.sep + args_output_file)
    else:
        # No output file specified - using default
        output_filename = os.path.normpath(
            convert_vars.BASE_PATH
            + os.sep
            + convert_vars.DEFAULT_OUTPUT_FILENAME
            + ("_template" if convert_vars.making_template else "")
            + "."
            + file_type.strip(".")
        )

    logging.debug(f" --- output_filename before fix extension = {output_filename}")
    output_filename = check_fix_file_extension(output_filename, file_type)
    logging.debug(f" --- output_filename AFTER fix extension = {output_filename}")

    # Do the replacement of filename place-holders with meta data
    find_replace = get_find_replace_list(meta)
    f = os.path.basename(output_filename)
    for r in find_replace:
        f = f.replace(*r)
    output_filename = os.path.dirname(output_filename) + os.sep + f

    logging.debug(f" --- output_filename = {output_filename}")
    return output_filename


def check_fix_file_extension(filename: str, file_type: str) -> str:
    if filename and not filename.endswith(file_type):
        filename_split = os.path.splitext(filename)
        if filename_split[1].strip(".").isnumeric():
            filename = filename + "." + file_type.strip(".")
        else:
            filename = ".".join([os.path.splitext(filename)[0], file_type.strip(".")])
        logging.debug(f" --- output_filename with new ext = {filename}")
    return filename


def get_find_replace_list(meta: Dict[str, str]) -> List[Tuple[str, str]]:
    ll: List[Tuple[str, str]] = [
        ("_type", "_" + meta["edition"].lower()),
        ("_edition", "_" + meta["edition"].lower()),
        ("_component", "_" + meta["component"].lower()),
        ("_language", "_" + meta["language"].lower()),
        ("_lang", "_" + meta["language"].lower()),
        ("_version", "_" + meta["version"].lower()),
        ("_ver", "_" + meta["version"].lower()),
    ]
    if not convert_vars.making_template:
        ll.append(("_template", ""))
    return ll


def get_meta_data(language_data: Dict[str, Dict[str, str]]) -> Dict[str, str]:
    meta = {}
    if "meta" in list(language_data.keys()):
        for key, value in language_data["meta"].items():
            if key in ("edition", "component", "language", "version"):
                meta[key] = value
        return meta
    else:
        logging.error(
            "Could not find meta tag in the language data. "
            "Please ensure required language file is in the source folder."
        )
    logging.debug(f" --- meta data = {meta}")
    return meta


def get_replacement_data(
    yaml_files: List[str], data_type: str = "translation", language: str = ""
) -> Dict[Any, Dict[Any, Any]]:
    """Get the raw data of the replacement text from correct yaml file"""
    data = {}
    logging.debug(f" --- Starting get_replacement_data() for data_type = {data_type} and language = {language}")
    if convert_vars.making_template:
        lang = "en"
    else:
        lang = language
    for file in yaml_files:
        if os.path.splitext(file)[1] in (".yaml", ".yml") and (
            os.path.basename(file).find("-" + lang + ".") >= 0
            or os.path.basename(file).find("-" + lang.replace("-", "_") + ".") >= 0
            or os.path.basename(file).find("mappings") >= 0
        ):
            with open(file, "r", encoding="utf-8") as f:
                try:
                    data = yaml.load(f, Loader=yaml.BaseLoader)
                except yaml.YAMLError as e:
                    logging.info(f"Error loading yaml file: {file}. Error = {e}")
                    continue

            if data_type in ("translation", "translations") and (
                (data["meta"]["language"].lower() == language)
                or (data["meta"]["language"].lower() == "en" and language == "template")
            ):
                logging.debug(" --- found source language file: " + os.path.split(file)[1])
                break
            elif (
                data_type in ("mapping", "mappings")
                and "meta" in data.keys()
                and "component" in data["meta"].keys()
                and data["meta"]["component"] == "mappings"
            ):
                logging.debug(" --- found mappings file: " + os.path.split(file)[1])
                break

            else:
                logging.debug(" --- found source file: " + os.path.split(file)[1])
                if "meta" in list(data.keys()):
                    meta_keys = data["meta"].keys()
                    logging.debug(f" --- data.keys() = {data.keys()}, data[meta].keys() = {meta_keys}")
                data = {}
                continue
    if not data or "suits" not in list(data.keys()):
        logging.error("Could not get language data from yaml files.")
    logging.debug(f" --- Len = {len(data)}.")
    return data


def get_replacement_dict(input_data: Dict[str, Any], mappings: bool = False) -> Dict[str, str]:
    """Loop through language file data and build up a find-replace dict"""
    data = {}
    for key in list(k for k in input_data.keys() if k != "meta"):
        suit_tags, suit_key = get_suit_tags_and_key(key)
        logging.debug(f" --- key = {key}.")
        logging.debug(f" --- suit_tags = {suit_tags}")
        logging.debug(f" --- suit_key = {suit_key}")

        for suit, suit_tag in zip(input_data[key], suit_tags):
            logging.debug(f" --- suit [name] = {suit['name']}")
            logging.debug(f" --- suit_tag = {suit_tag}")
            tag_for_suit_name = get_tag_for_suit_name(suit, suit_tag)
            data.update(tag_for_suit_name)

            card_tag = ""
            for card in suit[suit_key]:
                for tag, text_output in card.items():
                    if tag == "value":
                        continue

                    full_tag = get_full_tag(suit_tag, card["value"], tag)

                    # Add a translation for "Joker"
                    if suit_tag == "WC" and tag == "value":
                        full_tag = "${{{}}}".format("_".join([suit_tag, card_tag, tag]))

                    # Mappings is sometimes loaded as a list. Convert to string
                    if convert_vars.making_template:
                        text1 = check_make_list_into_text(text_output, False)
                        data[text1] = full_tag
                        if mappings:
                            data[text1.replace(", ", ",")] = full_tag
                            text1 = check_make_list_into_text(text_output, True)
                            data[text1] = full_tag
                    else:
                        data[full_tag] = check_make_list_into_text(text_output, True)
    if convert_vars.args.debug and not mappings:
        debug_txt = " --- Translation data showing First 4 (key: text):\n* "
        debug_txt += "\n* ".join(l1 + ": " + str(data[l1]) for l1 in list(data.keys())[:4])
        logging.debug(debug_txt)
        debug_txt = " --- Translation data showing Last 4 (key: text):\n* "
        debug_txt += "\n* ".join(l1 + ": " + str(data[l1]) for l1 in list(data.keys())[-4:])
        logging.debug(debug_txt)
    return data


def check_make_list_into_text(var: List[str], group_numbers: bool = True) -> str:
    if isinstance(var, list):
        if group_numbers:
            var = group_number_ranges(var)
        text_output = ", ".join(str(s) for s in list(var))
        if len(text_output.strip()) == 0:
            text_output = " - "
        return text_output
    else:
        return str(var)


def group_number_ranges(data: List[str]) -> List[str]:
    if len(data) < 2 or len([s for s in data if not s.isnumeric()]):
        return data
    list_ranges: List[str] = []
    data_numbers = [int(s) for s in data]
    for k, g in groupby(enumerate(data_numbers), lambda x: x[0] - x[1]):
        group: List[int] = list(map(itemgetter(1), g))
        group = list(map(int, group))
        if group[0] == group[-1]:
            list_ranges.append(str(group[0]))
        else:
            list_ranges.append(str(group[0]) + "-" + str(group[-1]))
    return list_ranges


def get_suit_tags_and_key(key: str) -> Tuple[List[str], str]:
    # Short tags to match the suits in the template documents
    suit_tags: List[str] = []
    suit_key: str = ""
    if key == "suits":
        suit_tags = ["VE", "AT", "SM", "AZ", "CR", "CO", "WC"]
        suit_key = "cards"
    elif key == "paragraphs":
        suit_tags = ["Common"]
        suit_key = "sentences"
    return suit_tags, suit_key


def get_tag_for_suit_name(suit: Dict[str, Any], suit_tag: str) -> Dict[str, str]:
    data: Dict[str, str] = {}
    logging.debug(f" --- suit_tag = {suit_tag}, suit[name] = {suit['name']}")
    if convert_vars.making_template:
        data[suit["name"]] = "${{{}}}".format(suit_tag + "_suit")
        if suit_tag == "WC":
            data["Joker"] = "${WC_Joker}"
    else:
        data["${{{}}}".format(suit_tag + "_suit")] = suit["name"]
        if suit_tag == "WC":
            data["${WC_Joker}"] = "Joker"
    logging.debug(f" --- making_template {convert_vars.making_template}. suit_tag dict = {data}")
    return data


def get_full_tag(suit_tag: str, card: str, tag: str) -> str:
    if suit_tag == "WC":
        full_tag = "${{{}}}".format("_".join([suit_tag, card, tag]))
    elif suit_tag == "Common":
        full_tag = "${{{}}}".format("_".join([suit_tag, card]))
    else:
        full_tag = "${{{}}}".format("_".join([suit_tag, suit_tag + card, tag]))
    return full_tag


def zip_dir(path: str, zip_filename: str) -> None:
    """Zip all the files recursively from path into zip_filename (excluding root path)"""
    with zipfile.ZipFile(zip_filename, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for root, dirs, files in os.walk(path):
            for file in files:
                f = os.path.join(root, file)
                zip_file.write(f, f[len(path) :])


def ensure_folder_exists(folder_path: str) -> None:
    """Check if folder exists and if not, create folders recursively."""
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)


def replace_docx_inline_text(doc: docx.Document, data: Dict[str, str]) -> docx.Document:
    """Replace the text in the docx document."""
    logging.debug(" --- starting docx_replace")

    if convert_vars.making_template:
        replacement_values = sort_keys_longest_to_shortest(data)
    else:
        replacement_values = list(data.items())

    paragraphs = get_document_paragraphs(doc)
    for p in paragraphs:
        runs_text = "".join(r.text for r in p.runs)
        if runs_text.strip() == "" or (
            convert_vars.making_template and re.search(re.escape("${") + ".*" + re.escape("}"), runs_text)
        ):
            continue
        for key, val in replacement_values:
            replaced_key = False
            for i, run in enumerate(p.runs):
                if run.text.find(key) != -1:
                    p.runs[i].text = run.text.replace(key, val)
                    replaced_key = True
                    runs_text = runs_text.replace(key, val)
            if not replaced_key:
                if runs_text.find(key) != -1:
                    runs_text = runs_text.replace(key, val)
                    for i, r in enumerate(p.runs):
                        p.runs[i].text = ""
                    p.runs[0].text = runs_text

    logging.debug(" --- finished replacing text in doc")
    return doc


def get_document_paragraphs(doc: docx) -> List[docx.Document]:
    paragraphs = list(doc.paragraphs)
    l1 = len(paragraphs)
    for table in doc.tables:
        paragraphs += get_paragraphs_from_table_in_doc(table)
    l2 = len(paragraphs)
    if not len(paragraphs):
        logging.error("No paragraphs found in doc")
    logging.debug(f" --- count doc paragraphs = {l1}, with table paragraphs = {l2}")
    return paragraphs


def get_paragraphs_from_table_in_doc(doc_table: docx.Document) -> List[docx.Document]:
    paragraphs: List[docx.Document] = []
    for row in doc_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if len(paragraph.runs):
                    paragraphs.append(paragraph)
            for t2 in cell.tables:
                paragraphs += get_paragraphs_from_table_in_doc(t2)
    return paragraphs


def get_docx_document(docx_file: str) -> docx.Document:
    """Open the file and return the docx document."""
    if os.path.isfile(docx_file):
        return docx.Document(docx_file)
    else:
        logging.error("Could not find file at: " + str(docx_file))
        return docx.Document()


def get_files_from_of_type(path: str, ext: str) -> List[str]:
    """Get a list of files from a specified folder recursively, that have the specified extension."""
    files = []
    for root, dirnames, filenames in os.walk(path):
        for filename in fnmatch.filter(filenames, "*." + str(ext)):
            files.append(os.path.join(root, filename))
    if not files:
        logging.error("No language files found in folder: " + str(os.sep.join([convert_vars.BASE_PATH, "source"])))
    logging.debug(f" --- found {len(files)} files of type {ext}. Showing first few:\n* " + str("\n* ".join(files[:3])))
    return files


def set_making_template() -> None:
    if hasattr(convert_vars.args, "language"):
        convert_vars.making_template = convert_vars.args.language.lower() == "template"
    else:
        convert_vars.making_template = False


def parse_arguments(input_args: List[str]) -> argparse.Namespace:
    """Parse and validate the input arguments. Return object containing argument values."""
    description = "Tool to output OWASP Cornucopia playing cards into different file types and languages. "
    description += "\nExample usage: $ ./cornucopia/convert.py -t docx -l es "
    description += "\nExample usage: c:\\cornucopia\\scripts\\convert.py -t idml -l fr "
    description += "-o 'my_output_folder/owasp_cornucopia_edition_language_version.idml'"
    parser = argparse.ArgumentParser(description=description, formatter_class=argparse.RawTextHelpFormatter)
    parser.add_argument(
        "-i",
        "--inputfile",
        type=str,
        default="",
        help=(
            "Input (template) file to use."
            f"\nDefault={convert_vars.DEFAULT_TEMPLATE_FILENAME}.(docx|idml)"
            "\nTemplate type is dependent on output type (-t) or file (-o) specified."
        ),
    )
    group = parser.add_mutually_exclusive_group(required=False)
    group.add_argument(
        "-t",
        "--outputfiletype",
        type=str,
        choices=convert_vars.FILETYPE_CHOICES,
        help="Type of file to output. Default = docx. If specified, this overwrites the output file extension",
    )
    parser.add_argument(
        "-o",
        "--outputfile",
        default="",
        type=str,
        help=(
            "Specify a path and name of output file to generate. (caution: existing file will be overwritten). "
            f"\ndefault = {convert_vars.DEFAULT_OUTPUT_FILENAME}.(docx|pdf|idml)"
        ),
    )
    group = parser.add_mutually_exclusive_group(required=False)
    group.add_argument(
        # parser.add_argument(
        "-l",
        "--language",
        type=str,
        choices=convert_vars.LANGUAGE_CHOICES,
        default="en",
        help=(
            "Output language to produce. [`en`, `es`, `fr`, `pt-br`, `template`] "
            "\nTemplate will attempt to create a template from the english input file and "
            "\nreplacing strings with the template lookup codes"
        ),
    )
    parser.add_argument(
        "-d",
        "--debug",
        action="store_true",
        help="Output additional information to debug script",
    )
    args = parser.parse_args(input_args)
    return args


if __name__ == "__main__":
    convert_vars: ConvertVars = ConvertVars()
    main()
